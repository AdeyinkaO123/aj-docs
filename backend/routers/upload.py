import io
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from database import get_db
from models import Document, User
from schemas import DocumentOut
import markdown

router = APIRouter(prefix="/api/upload", tags=["upload"])

MAX_SIZE = 5 * 1024 * 1024  # 5MB


@router.post("/", response_model=DocumentOut)
async def upload_file(
    user_id: int = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
):
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    content_bytes = await file.read()
    if len(content_bytes) > MAX_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 5MB)")

    filename = file.filename or ""
    if "." in filename:
        ext = filename.rsplit(".", 1)[-1].lower()
        title = filename.rsplit(".", 1)[0]
    else:
        ext = ""
        title = filename or "Uploaded Document"

    if ext == "txt":
        text = content_bytes.decode("utf-8", errors="replace")
        html = "".join(f"<p>{line}</p>" for line in text.splitlines() if line.strip()) or "<p></p>"
    elif ext == "md":
        text = content_bytes.decode("utf-8", errors="replace")
        html = markdown.markdown(text, extensions=["extra"])
    elif ext == "docx":
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx not installed")

        doc_obj = DocxDocument(io.BytesIO(content_bytes))
        parts = []
        for para in doc_obj.paragraphs:
            style = para.style.name if para.style else ""
            if style.startswith("Heading 1"):
                tag = "h1"
            elif style.startswith("Heading 2"):
                tag = "h2"
            elif style.startswith("Heading 3"):
                tag = "h3"
            else:
                tag = "p"

            inner = ""
            for run in para.runs:
                text = run.text
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                if run.underline:
                    text = f"<u>{text}</u>"
                inner += text

            if inner.strip():
                parts.append(f"<{tag}>{inner}</{tag}>")

        html = "".join(parts) or "<p></p>"
    else:
        raise HTTPException(status_code=422, detail=f"Unsupported file type: .{ext or '(none)'}")

    doc = Document(title=title, content=html, owner_id=user_id)
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return {**doc.__dict__, "owner": doc.owner, "is_owner": True, "is_shared": False}
